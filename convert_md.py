import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_run_docx():
    with open("run.md", "r", encoding="utf-8") as f:
        md_text = f.read()

    # Very simple markdown to docx mapping
    doc = Document()
    
    # Simple parsing line by line since it's a simple README
    lines = md_text.split('\n')
    
    in_code_block = False
    code_content = ""
    
    for line in lines:
        if line.startswith('```'):
            if in_code_block:
                # End code block
                p = doc.add_paragraph()
                p.style = 'Normal'
                run = p.add_run(code_content.strip())
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                # Background gray hack: no direct shading in simple docx, just change color
                run.font.color.rgb = RGBColor(50, 50, 50)
                code_content = ""
                in_code_block = False
            else:
                in_code_block = True
            continue
            
        if in_code_block:
            code_content += line + "\n"
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.strip() == '---':
            pass # skip hr
        elif line.strip() == '':
            pass # skip blank
        else:
            # check for bold
            p = doc.add_paragraph()
            import re
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                else:
                    p.add_run(part)

    doc.save("run.docx")

if __name__ == "__main__":
    create_run_docx()
